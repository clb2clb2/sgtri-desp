#!/usr/bin/env python3
"""
Generador de `docs/Estructura_Input_Desplazamiento.docx`.
Requiere: python-docx

Ejecución (PowerShell):
  .\.venv\Scripts\Activate.ps1
  C:/Proyectos-Github/sgtri-desp/.venv/Scripts/python.exe tools\generate_input_schema_docx.py
"""
from docx import Document
from docx.shared import Pt

schema = [
    ("fechaIda", "string", "Obligatorio", "Fecha de salida en formato dd/mm/aa o dd/mm/yyyy. Ej: 01/12/2025"),
    ("horaIda", "string", "Requerido salvo en _segmentMode", "Hora de salida hh:mm. Ej: 09:00"),
    ("fechaRegreso", "string", "Obligatorio", "Fecha de regreso dd/mm/aa. Ej: 03/12/2025"),
    ("horaRegreso", "string", "Requerido salvo en _segmentMode", "Hora de regreso hh:mm. Ej: 20:00"),
    ("cruceIda", "string", "Opcional", "Fecha del cruce de frontera de ida (internacional). Formato dd/mm/aa."),
    ("cruceVuelta", "string", "Opcional", "Fecha del cruce de frontera de vuelta (internacional)."),
    ("pais", "string", "Opcional", "Nombre del país destino (ej. 'Francia')."),
    ("paisIndex", "number", "Opcional (preferible)", "Índice en `dietasPorPais.paises` (0 = España)."),
    ("km", "number|string", "Opcional", "Kilómetros totales. Acepta números o text como '123,4'."),
    ("alojamiento", "number|string", "Opcional", "Importe total alojamiento en euros."),
    ("ticketCena", "boolean", "Opcional", "Si aporta justificante de la cena del último día."),
    ("tipoProyecto", "string", "Opcional", "Clave del tipo de proyecto; usada para decidir normativa (RD vs Decreto)."),
    ("kmTarifa", "number", "Opcional", "Tarifa por km. Si no se provee se usa valor por defecto o `window.__sgtriDatos.kmTarifas`)."),
    ("residenciaEventual", "boolean", "Opcional", "Aplica coeficiente (ej. 0.8) para cálculo de IRPF si true."),
    ("excludeManutencion", "boolean", "Opcional", "Si true, excluir manutención del cálculo."),
    ("justificarPernocta", "boolean", "Opcional", "Si true, el último tramo puede aumentar noches (equivalente a justificar pernocta)."),
    ("excludeAlojamiento", "boolean", "Opcional", "Si true, excluir alojamiento del cálculo."),
    ("_segmentMode", "boolean", "Opcional (interno)", "Usado por el wrapper para cálculos parciales de tramo; normalmente no lo establece el usuario.")
]

example = {
    "fechaIda": "01/12/2025",
    "horaIda": "09:00",
    "fechaRegreso": "03/12/2025",
    "horaRegreso": "20:00",
    "pais": "Francia",
    "paisIndex": 1,
    "km": "0",
    "alojamiento": "150.00",
    "ticketCena": False,
    "tipoProyecto": "PROY-XYZ",
    "kmTarifa": 0.26,
    "residenciaEventual": False,
    "excludeManutencion": False,
    "justificarPernocta": False,
    "excludeAlojamiento": False
}

def create_docx(path):
    doc = Document()
    h = doc.add_heading('Estructura del objeto input para calculateDesplazamiento', level=1)
    p = doc.add_paragraph()
    p.add_run('Descripción').bold = True
    doc.add_paragraph('Este documento describe todos los campos esperados en el objeto `input` que se pasa a la función `calculateDesplazamiento(input)`.')

    doc.add_paragraph()
    doc.add_paragraph('Esquema de campos').bold = True
    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Campo'
    hdr_cells[1].text = 'Tipo'
    hdr_cells[2].text = 'Obligatorio'
    hdr_cells[3].text = 'Descripción / Ejemplo'

    for fld, typ, req, desc in schema:
        row_cells = table.add_row().cells
        row_cells[0].text = fld
        row_cells[1].text = typ
        row_cells[2].text = req
        row_cells[3].text = desc

    doc.add_paragraph()
    doc.add_paragraph('Ejemplo de objeto `input` (JSON)').bold = True
    doc.add_paragraph(str(example))

    doc.add_paragraph()
    doc.add_paragraph('Notas técnicas').bold = True
    doc.add_paragraph(' - La función es independiente del DOM: sólo usa los campos del objeto `input`.')
    doc.add_paragraph(' - Para datos auxiliares (tablas de dietas, límites IRPF, tarifas km, normativas) el motor consulta `window.__sgtriDatos` si no se inyectan en `input` mediante `input.__sgtriDatos`.')
    doc.add_paragraph(' - El wrapper en `calculoDesp.js` puede construir entradas parciales para tramos con `_segmentMode = true` (campo top-level).' )

    doc.save(path)

if __name__ == '__main__':
    import os
    outdir = os.path.join(os.path.dirname(__file__), '..', 'docs')
    os.makedirs(outdir, exist_ok=True)
    outpath = os.path.join(outdir, 'Estructura_Input_Desplazamiento.docx')
    create_docx(outpath)
    print('Generado:', outpath)
