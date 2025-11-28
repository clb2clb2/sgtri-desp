#!/usr/bin/env python3
"""
Genera `docs/formLogic.docx` describiendo las secciones y recomendaciones de refactor
para `js/formLogic.js`.
Requisitos: python-docx
Ejecución (PowerShell):
  pip install python-docx
  python tools\generate_formLogic_docx.py
"""
import re
from docx import Document
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
JS_PATH = ROOT / 'js' / 'formLogic.js'
OUT_DIR = ROOT / 'docs'
OUT_DIR.mkdir(exist_ok=True)
OUT_PATH = OUT_DIR / 'formLogic.docx'

text_intro = (
    "Este documento describe la estructura de `js/formLogic.js`, sus dependencias, "
    "las secciones funcionales principales y recomendaciones de refactorización.\n\n"
    "El objetivo es facilitar dividir el archivo en módulos más pequeños y testables, "
    "mover la lógica de cálculo al motor puro y delegar la renderización a `salidaDesp`.")

# Leer archivo JS
src = JS_PATH.read_text(encoding='utf-8')

# Encontrar funciones top-level
fnames = re.findall(r'function\s+([A-Za-z0-9_]+)\s*\(', src)
# Encontrar const/let arrow functions assigned to names: const name = ( ... ) => { }
fnames += re.findall(r'const\s+([A-Za-z0-9_]+)\s*=\s*\(', src)
fnames = list(dict.fromkeys(fnames))

# Buscar secciones por palabras clave
has_date_parsers = 'parseDateStrict' in src or 'parseTimeStrict' in src
has_grouped_input = 'grouped' in src or 'vincular' in src
has_create_ficha = 'create' in src and 'ficha' in src

# Recomendar módulos
recom = [
    ('Sanitizers y utilidades de parseo',
     'Funciones que normalizan números/monedas y parsean fechas/horas. Mover a `js/utils/parse.js` o `js/cogeDatosDesp.js`.'),
    ('Validación de fecha/hora',
     'Funciones que validan pares fecha/hora y marcan campos inválidos. Extraer a `js/validators/datetime.js` y exponer funciones puras que devuelvan errores y marcas.'),
    ('Gestión de fichas (crear/eliminar)',
     'DOM builders y lógica de añadir/quitar desplazamientos. Mantener aquí sólo la manipulación de DOM y plantilla; extraer la lógica de nombres/IDs y estructura a un helper.'),
    ('Listeners y delegación de eventos',
     'Agrupar registradores de eventos y usar delegación cuando sea posible. Mover la programación/debounce a `js/logicaDesp.js`.'),
    ('Otros gastos (líneas dinámicas)',
     'Extraer la creación/serialización de líneas de "otros gastos" en un módulo propio para facilitar tests y reutilización.'),
    ('Integración con el motor',
     'El archivo debe delegar cálculos a `window.calculoDesp` y solo encargarse de construir el `calcInput` y montar el resultado. Evitar cálculos aritméticos en el DOM.'),
]

# Crear documento
doc = Document()

doc.add_heading('Documentación de js/formLogic.js', level=1)

doc.add_paragraph(text_intro)

doc.add_heading('Resumen rápido', level=2)
p = doc.add_paragraph()
p.add_run('Funciones principales detectadas: ').bold = True
p.add_run(', '.join(fnames[:12]) + (', ...' if len(fnames) > 12 else ''))

if has_date_parsers:
    doc.add_paragraph('Contiene parsers y validadores de fecha/hora (p. ej. parseDateStrict/parseTimeStrict).')

# Dependencias globales
doc.add_heading('Dependencias globales', level=2)
doc.add_paragraph('El archivo usa o asume disponibilidad de las siguientes entidades globales:')
for g in ['window.calculoDesp', 'window.salidaDesp', 'scheduleFullRecalc / logicaDesp', 'document (DOM)']:
    doc.add_paragraph('- ' + g)

# Secciones detectadas
doc.add_heading('Secciones funcionales', level=2)
sections = [
    ('Sanitizers y parseo', 'Normalización de inputs de números y fechas; parseNumber, formatters.'),
    ('Creación/gestión de fichas', 'Funciones para crear, clonar y eliminar grupos `.desplazamiento-grupo` en el DOM.'),
    ('Attach/registradores de eventos', 'Funciones que añaden listeners a inputs (blur, change) y delegan recalculos.'),
    ('Validación', 'validateDateTimePairAndUpdateUI, validateCrucesAndUpdateUI y similares.'),
    ('Helpers de UI (tooltip/warn)', 'ensureGlobalWarnTooltip, attachWarnHandlers y utilidades para tooltips).'),
    ('Integración con el motor', 'Llamados a `window.calculoDesp.calculaDesplazamientoFicha` y funciones auxiliares como computeDescuentoManutencion.'),
]
for t,b in sections:
    doc.add_paragraph().add_run(t + ':').bold = True
    doc.add_paragraph('  ' + b)

# Recomendaciones de refactor
doc.add_heading('Recomendaciones de refactor', level=2)
for title, body in recom:
    doc.add_paragraph().add_run(title + ':').bold = True
    doc.add_paragraph('  ' + body)

# Checklist para descomposición
doc.add_heading('Checklist para descomposición', level=2)
checks = [
    'Extraer funciones puras de parseo/normalización a un módulo utilitario.',
    'Extraer validadores de fecha/hora y pruebas unitarias para ellos.',
    'Mantener creación de DOM y plantillas en un módulo de UI separado (p. ej. js/ui/desplazamientos.js).',
    'Reemplazar manipulación directa de clases CSS por funciones utilitarias (addClass/removeClass).',
    'Usar delegación de eventos para inputs dinámicos en lugar de listeners individuales cuando sea posible.',
    'Evitar cálculos monetarios en el DOM; delegar al motor `calculateDesplazamiento` y usar `renderSalidaHtml` para generar output.'
]
for c in checks:
    doc.add_paragraph('- ' + c)

# Incluir resumen de funciones encontradas (primeras 40)
doc.add_heading('Funciones detectadas (ejemplo)', level=2)
for i,name in enumerate(fnames[:40], start=1):
    doc.add_paragraph(f'{i}. {name}')

# Añadir nota final
doc.add_paragraph()
doc.add_paragraph('Sugerencia: crear un PR pequeño por cada módulo extraído para facilitar revisión.')

# Guardar
doc.save(str(OUT_PATH))
print('Generado', OUT_PATH)
