#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Genera el documento Word "Lógica interna del cálculo de desplazamientos"
Documentación técnica para revisores del código.
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_shading(cell, color):
    """Aplica color de fondo a una celda."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_code_block(doc, code_text):
    """Añade un bloque de código con formato monoespaciado."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(40, 40, 40)
    return p

def create_document():
    doc = Document()
    
    # Configurar márgenes
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # =========================================================================
    # TÍTULO
    # =========================================================================
    title = doc.add_heading('Lógica Interna del Cálculo de Desplazamientos', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Documentación técnica para revisores del código')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.color.rgb = RGBColor(80, 80, 80)
    
    doc.add_paragraph()
    
    # =========================================================================
    # ÍNDICE
    # =========================================================================
    doc.add_heading('Índice', level=1)
    
    indice = [
        '1. Introducción y propósito',
        '2. Arquitectura general',
        '3. Módulos del sistema',
        '4. Flujo de datos paso a paso',
        '5. Reglas de negocio aplicadas',
        '6. Estructura de datos',
        '7. Casos especiales',
        '8. Puntos clave para modificar la lógica'
    ]
    
    for item in indice:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Cm(1)
    
    doc.add_page_break()
    
    # =========================================================================
    # 1. INTRODUCCIÓN
    # =========================================================================
    doc.add_heading('1. Introducción y propósito', level=1)
    
    p = doc.add_paragraph()
    p.add_run('Este documento explica cómo funciona internamente el sistema de cálculo de desplazamientos. ')
    p.add_run('Está dirigido a personas que necesiten:')
    
    bullets = [
        'Entender cómo se calculan los importes',
        'Localizar dónde se aplica una regla de negocio específica',
        'Modificar alguna parte del cálculo'
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('Importante: ').bold = True
    p.add_run('No es necesario ser programador experto para entender este documento, pero sí ayuda tener nociones básicas de JavaScript y HTML.')
    
    doc.add_paragraph()
    
    # =========================================================================
    # 2. ARQUITECTURA GENERAL
    # =========================================================================
    doc.add_heading('2. Arquitectura general', level=1)
    
    doc.add_heading('2.1. Visión de alto nivel', level=2)
    
    p = doc.add_paragraph()
    p.add_run('El sistema sigue un flujo lineal de tres fases:')
    
    # Diagrama de flujo simple con texto
    diagram = doc.add_paragraph()
    diagram.alignment = WD_ALIGN_PARAGRAPH.CENTER
    diagram.add_run('┌─────────────┐      ┌─────────────┐      ┌─────────────┐\n').font.name = 'Consolas'
    diagram.add_run('│  RECOGER    │  →   │  CALCULAR   │  →   │  MOSTRAR    │\n').font.name = 'Consolas'
    diagram.add_run('│   DATOS     │      │  IMPORTES   │      │ RESULTADO   │\n').font.name = 'Consolas'
    diagram.add_run('└─────────────┘      └─────────────┘      └─────────────┘\n').font.name = 'Consolas'
    diagram.add_run('  cogeDatosDesp        calculoDesp          salidaDesp').font.name = 'Consolas'
    
    p = doc.add_paragraph()
    p.add_run('Cada fase tiene un módulo JavaScript dedicado. ')
    p.add_run('Esta separación permite modificar una parte sin afectar a las demás.')
    
    doc.add_paragraph()
    
    doc.add_heading('2.2. Principio de separación', level=2)
    
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    
    hdr = table.rows[0].cells
    hdr[0].text = 'Fase'
    hdr[1].text = 'Responsabilidad'
    hdr[2].text = 'Módulo'
    for cell in hdr:
        set_cell_shading(cell, 'D9EAD3')
        cell.paragraphs[0].runs[0].bold = True
    
    data = [
        ('Entrada', 'Leer datos del formulario y validarlos', 'cogeDatosDesp.js'),
        ('Proceso', 'Aplicar reglas y calcular importes', 'calculoDesp.js'),
        ('Salida', 'Generar HTML con el resultado', 'salidaDesp.js')
    ]
    
    for i, (fase, resp, modulo) in enumerate(data, 1):
        row = table.rows[i].cells
        row[0].text = fase
        row[1].text = resp
        row[2].text = modulo
    
    doc.add_paragraph()
    
    # =========================================================================
    # 3. MÓDULOS DEL SISTEMA
    # =========================================================================
    doc.add_heading('3. Módulos del sistema', level=1)
    
    doc.add_heading('3.1. cogeDatosDesp.js - Recolección de datos', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Propósito: ').bold = True
    p.add_run('Extraer todos los valores introducidos por el usuario en una ficha de desplazamiento y normalizarlos.')
    
    p = doc.add_paragraph()
    p.add_run('Función principal: ').bold = True
    p.add_run('collectDataFromFicha(despEl)')
    
    p = doc.add_paragraph('Qué hace:')
    
    bullets = [
        'Lee los campos del formulario (fechas, horas, km, alojamiento, país...)',
        'Parsea las fechas de formato "dd/mm/aa" a objetos Date de JavaScript',
        'Convierte los importes de texto (ej: "1.500 km") a números',
        'Detecta si es un viaje internacional (país ≠ España)',
        'Valida la coherencia de fechas (ida < regreso, cruces correctos)',
        'Devuelve un objeto normalizado con todos los datos listos para calcular'
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('Ubicación de las validaciones: ').bold = True
    p.add_run('Función validarFechas() en este mismo archivo.')
    
    doc.add_paragraph()
    
    doc.add_heading('3.2. calculoDesp.js - Motor de cálculo', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Propósito: ').bold = True
    p.add_run('Aplicar todas las reglas de negocio y calcular los importes finales.')
    
    p = doc.add_paragraph()
    p.add_run('Este es el módulo más extenso porque contiene toda la lógica de cálculo. ')
    p.add_run('Está dividido en dos secciones:')
    
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    
    hdr = table.rows[0].cells
    hdr[0].text = 'Sección'
    hdr[1].text = 'Descripción'
    hdr[2].text = 'Funciones clave'
    for cell in hdr:
        set_cell_shading(cell, 'CFE2F3')
        cell.paragraphs[0].runs[0].bold = True
    
    table.rows[1].cells[0].text = 'Motor (puro)'
    table.rows[1].cells[1].text = 'Funciones matemáticas que no tocan el DOM'
    table.rows[1].cells[2].text = 'calculateDesplazamiento(), calcManutenciones(), calcNoches()'
    
    table.rows[2].cells[0].text = 'Wrapper (orquestador)'
    table.rows[2].cells[1].text = 'Conecta DOM → Motor → Renderizador'
    table.rows[2].cells[2].text = 'calculaDesplazamientoFicha(), buildSegmentInputs()'
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('Función principal del motor: ').bold = True
    p.add_run('calculateDesplazamiento(input)')
    
    p = doc.add_paragraph('Recibe un objeto con los datos de entrada y devuelve:')
    
    bullets = [
        'Número de manutenciones y su importe',
        'Número de noches y su importe máximo',
        'Importe de kilometraje',
        'Información sobre IRPF sujeto a retención',
        'Flags de ambigüedad (ej: última noche dudosa)'
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')
    
    doc.add_paragraph()
    
    doc.add_heading('3.3. salidaDesp.js - Renderizado', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Propósito: ').bold = True
    p.add_run('Convertir los números calculados en HTML visual que el usuario pueda ver.')
    
    p = doc.add_paragraph()
    p.add_run('Función principal: ').bold = True
    p.add_run('renderSalida(despEl, salidaData)')
    
    p = doc.add_paragraph('Estructura interna:')
    
    bullets = [
        'templates: Funciones que generan fragmentos HTML (línea de manutención, línea de km, etc.)',
        'renderSimple(): Para desplazamientos nacionales (sin segmentos)',
        'renderSegmentado(): Para desplazamientos internacionales (con tramos España/Extranjero)',
        'mountSalida(): Inserta el HTML generado en la página'
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')
    
    doc.add_paragraph()
    
    doc.add_heading('3.4. logicaDesp.js - Control de UI', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Propósito: ').bold = True
    p.add_run('Gestionar los eventos del formulario y coordinar cuándo se disparan los recálculos.')
    
    p = doc.add_paragraph('Qué hace:')
    
    bullets = [
        'Escucha cambios en los campos del formulario (blur, change)',
        'Aplica un "debounce" para no recalcular en cada tecla',
        'Muestra/oculta campos condicionales (fronteras para viajes internacionales, ticket cena)',
        'Valida visualmente los campos (marca errores en rojo)',
        'Llama a calculaDesplazamientoFicha() cuando hay cambios'
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')
    
    doc.add_paragraph()
    
    doc.add_heading('3.5. Otros módulos auxiliares', level=2)
    
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    
    hdr = table.rows[0].cells
    hdr[0].text = 'Módulo'
    hdr[1].text = 'Función'
    for cell in hdr:
        set_cell_shading(cell, 'F4CCCC')
        cell.paragraphs[0].runs[0].bold = True
    
    data = [
        ('limpiaDatos.js', 'Sanitizar y formatear datos de entrada (fechas, números, texto)'),
        ('formLogic.js', 'Lógica general del formulario (no específica de desplazamientos)'),
        ('datos.json', 'Tablas de precios por país, tarifas km, límites IRPF')
    ]
    
    for i, (mod, func) in enumerate(data, 1):
        table.rows[i].cells[0].text = mod
        table.rows[i].cells[1].text = func
    
    doc.add_page_break()
    
    # =========================================================================
    # 4. FLUJO DE DATOS PASO A PASO
    # =========================================================================
    doc.add_heading('4. Flujo de datos paso a paso', level=1)
    
    p = doc.add_paragraph()
    p.add_run('Veamos qué ocurre desde que el usuario modifica un campo hasta que ve el resultado:')
    
    doc.add_heading('Paso 1: Evento de cambio', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Cuando el usuario sale de un campo (blur) o cambia un selector, ')
    p.add_run('logicaDesp.js').bold = True
    p.add_run(' captura el evento y programa un recálculo.')
    
    add_code_block(doc, 'Usuario modifica campo → logicaDesp detecta → scheduleRecalcForId(id)')
    
    doc.add_paragraph()
    
    doc.add_heading('Paso 2: Recolección de datos', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Se llama a ')
    p.add_run('cogeDatosDesp.collectDataFromFicha()').bold = True
    p.add_run(' que extrae todos los valores del formulario:')
    
    add_code_block(doc, '''Formulario HTML
    ↓
collectDataFromFicha(despEl)
    ↓
Objeto normalizado: {
    fechaIda: Date,
    fechaRegreso: Date,
    horaIda: {hh, mm},
    horaRegreso: {hh, mm},
    km: 240,
    alojamiento: 150.50,
    pais: "Francia",
    paisIndex: 35,
    esInternacional: true,
    tipoProyecto: "PEI",
    ...
}''')
    
    doc.add_paragraph()
    
    doc.add_heading('Paso 3: Construcción del input de cálculo', level=2)
    
    p = doc.add_paragraph()
    p.add_run('El wrapper de ')
    p.add_run('calculoDesp.js').bold = True
    p.add_run(' transforma los datos normalizados en el formato que necesita el motor:')
    
    add_code_block(doc, '''const calcInput = {
    fechaIda: "15/01/25",
    horaIda: "08:00",
    fechaRegreso: "17/01/25",
    horaRegreso: "19:00",
    pais: "Francia",
    paisIndex: 35,
    km: 240,
    alojamiento: 150.50,
    tipoProyecto: "PEI",
    kmTarifa: 0.26,
    excludeManutencion: false,
    justificarPernocta: false,
    ...
}''')
    
    doc.add_paragraph()
    
    doc.add_heading('Paso 4: Cálculo por el motor', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Se llama a ')
    p.add_run('calculateDesplazamiento(calcInput)').bold = True
    p.add_run(' que ejecuta secuencialmente:')
    
    bullets = [
        'Determinar la normativa (Decreto o RD) según tipo de proyecto',
        'Obtener precios de tablas según país y normativa',
        'Calcular manutenciones según horas de salida/regreso',
        'Calcular noches de alojamiento',
        'Calcular kilometraje',
        'Calcular IRPF sujeto a retención',
        'Aplicar factor de residencia eventual si procede (×0.8)'
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')
    
    doc.add_paragraph()
    
    doc.add_heading('Paso 5: Viajes internacionales - Segmentación', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Si el viaje es internacional, se divide en ')
    p.add_run('tramos').bold = True
    p.add_run(':')
    
    add_code_block(doc, '''Viaje internacional:
┌─────────────────┐
│ Tramo España    │  fechaIda → cruceIda
│ (ida)           │  Precios: España
├─────────────────┤
│ Tramo Extranjero│  cruceIda → cruceVuelta
│                 │  Precios: País destino
├─────────────────┤
│ Tramo España    │  cruceVuelta → fechaRegreso
│ (vuelta)        │  Precios: España
└─────────────────┘''')
    
    p = doc.add_paragraph()
    p.add_run('Cada tramo se calcula por separado y luego se suman los totales.')
    
    doc.add_paragraph()
    
    doc.add_heading('Paso 6: Construcción de datos de salida', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Se construye un objeto ')
    p.add_run('salidaData').bold = True
    p.add_run(' con toda la información necesaria para el renderizado:')
    
    add_code_block(doc, '''salidaData = {
    id: "1",
    totales: {
        manutencion: 133.35,
        alojamientoMax: 307.68,
        alojamientoUser: 150.50,
        km: 62.40,
        otrosGastos: 0,
        total: 346.25,
        irpfSujeto: 26.67
    },
    detalles: {
        manutenciones: 2.5,
        precioManutencion: 53.34,
        noches: 2,
        precioNoche: 153.84,
        km: 240,
        precioKm: 0.26
    },
    ui: {
        esInternacional: false,
        alojamientoExcedeMax: false,
        nochesAmbiguas: false,
        residenciaEventual: false
    }
}''')
    
    doc.add_paragraph()
    
    doc.add_heading('Paso 7: Renderizado HTML', level=2)
    
    p = doc.add_paragraph()
    p.add_run('salidaDesp.renderSalida()').bold = True
    p.add_run(' genera el HTML usando los templates:')
    
    add_code_block(doc, '''salidaData
    ↓
renderSimple() o renderSegmentado()
    ↓
HTML generado:
<div class="calc-result">
    <div class="calc-line">
        <span class="label">Manutención: 2,5 × 53,34 €</span>
        <span class="leader">...</span>
        <span class="amount">133,35 €</span>
    </div>
    ...
    <div class="calc-total">
        <span class="label">Total:</span>
        <span class="amount"><strong>346,25 €</strong></span>
    </div>
</div>''')
    
    doc.add_paragraph()
    
    doc.add_heading('Paso 8: Inserción en el DOM', level=2)
    
    p = doc.add_paragraph()
    p.add_run('mountSalida()').bold = True
    p.add_run(' inserta el HTML generado dentro de la ficha de desplazamiento, ')
    p.add_run('reemplazando cualquier resultado anterior.')
    
    doc.add_page_break()
    
    # =========================================================================
    # 5. REGLAS DE NEGOCIO
    # =========================================================================
    doc.add_heading('5. Reglas de negocio aplicadas', level=1)
    
    doc.add_heading('5.1. Determinación de la normativa', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Ubicación: ').bold = True
    p.add_run('calculoDesp.js → getNormativa()')
    
    add_code_block(doc, '''Tipo de proyecto → Normativa
────────────────────────────
G24, PEI, NAL    → R.D. 462/2002
UEX, PCO, JEX    → Decreto 42/2025''')
    
    doc.add_paragraph()
    
    doc.add_heading('5.2. Obtención de precios', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Ubicación: ').bold = True
    p.add_run('calculoDesp.js → getPrecios()')
    
    p = doc.add_paragraph('Los precios se leen del archivo ')
    p.add_run('datos.json').bold = True
    p.add_run(' según:')
    
    bullets = [
        'Índice del país (0 = España, 1 = Alemania, 2 = Andorra...)',
        'Normativa aplicable (decreto42_2025 o rd462_2002)',
        'Tipo de importe (manutencion o alojamiento)'
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')
    
    doc.add_paragraph()
    
    doc.add_heading('5.3. Cálculo de manutenciones', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Ubicación: ').bold = True
    p.add_run('calculoDesp.js → calcManutenciones(), calcManutencionesSameDay(), calcManutencionesSeveralDays()')
    
    p = doc.add_paragraph()
    p.add_run('Horas clave definidas como constantes:')
    
    add_code_block(doc, '''HORA_COMIDA     = 14:00  (840 minutos)
HORA_FIN_COMIDA = 16:00  (960 minutos)
HORA_CENA       = 22:00  (1320 minutos)''')
    
    p = doc.add_paragraph()
    p.add_run('Regla para día de salida:')
    
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    
    hdr = table.rows[0].cells
    hdr[0].text = 'Condición'
    hdr[1].text = 'Manutención'
    for cell in hdr:
        set_cell_shading(cell, 'FCE5CD')
        cell.paragraphs[0].runs[0].bold = True
    
    table.rows[1].cells[0].text = 'Sale antes de las 14:00'
    table.rows[1].cells[1].text = '1 completa'
    table.rows[2].cells[0].text = 'Sale entre 14:00 y 22:00'
    table.rows[2].cells[1].text = '0,5 (cena)'
    table.rows[3].cells[0].text = 'Sale después de las 22:00'
    table.rows[3].cells[1].text = '0'
    
    doc.add_paragraph()
    
    doc.add_heading('5.4. Cálculo de noches', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Ubicación: ').bold = True
    p.add_run('calculoDesp.js → calcNoches(), evalLastNightByHour()')
    
    p = doc.add_paragraph()
    p.add_run('Horas clave para pernocta:')
    
    add_code_block(doc, '''HORA_PERNOCTA_MIN = 01:00  →  NO pernocta
HORA_PERNOCTA_MAX = 07:00  →  SÍ pernocta
Entre ambas                →  Zona AMBIGUA (requiere justificación)''')
    
    doc.add_paragraph()
    
    doc.add_heading('5.5. Residencia eventual', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Ubicación: ').bold = True
    p.add_run('calculoDesp.js → isResidenciaEventual(), monthsBetween()')
    
    add_code_block(doc, '''España:     > 1 mes    → Factor 0.8 (80%)
Extranjero: > 3 meses  → Factor 0.8 (80%)''')
    
    doc.add_paragraph()
    
    doc.add_heading('5.6. IRPF sujeto a retención', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Ubicación: ').bold = True
    p.add_run('calculoDesp.js → calcIRPF(), getLimitesIRPF()')
    
    p = doc.add_paragraph('Los límites exentos se leen de datos.json → limitesIRPF:')
    
    add_code_block(doc, '''España:     [26.67, 53.34]  (sin pernocta, con pernocta)
Extranjero: [48.08, 91.35]  (sin pernocta, con pernocta)

IRPF_día = máx(0, manutención_día - límite_exento)''')
    
    doc.add_page_break()
    
    # =========================================================================
    # 6. ESTRUCTURA DE DATOS
    # =========================================================================
    doc.add_heading('6. Estructura de datos', level=1)
    
    doc.add_heading('6.1. datos.json - Tablas de configuración', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Este archivo contiene todas las tablas de referencia:')
    
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    
    hdr = table.rows[0].cells
    hdr[0].text = 'Sección'
    hdr[1].text = 'Contenido'
    for cell in hdr:
        set_cell_shading(cell, 'D9EAD3')
        cell.paragraphs[0].runs[0].bold = True
    
    data = [
        ('tiposProyecto', 'Lista de tipos de proyecto con sus códigos'),
        ('kmTarifas', 'Tarifas por km: coche (0.26), motocicleta (0.106)'),
        ('dietasPorPais', 'Precios de manutención y alojamiento por país'),
        ('limitesIRPF', 'Límites exentos para España y extranjero'),
        ('normativasPorTipoProyecto', 'Qué normativa aplica a cada tipo')
    ]
    
    for i, (sec, cont) in enumerate(data, 1):
        table.rows[i].cells[0].text = sec
        table.rows[i].cells[1].text = cont
    
    doc.add_paragraph()
    
    doc.add_heading('6.2. Estructura dietasPorPais', level=2)
    
    add_code_block(doc, '''{
  "paises": ["España", "Alemania", "Andorra", ...],  // índice 0, 1, 2...
  "decreto42_2025": {
    "alojamiento": [102.56, 155.66, 54.69, ...],     // mismo orden
    "manutencion": [53.34, 91.35, 91.35, ...]
  },
  "rd462_2002": {
    "alojamiento": [65.97, 132.82, 46.88, ...],
    "manutencion": [37.40, 59.50, 37.86, ...]
  }
}

Para obtener precio de manutención de Francia (índice 35) con Decreto:
→ dietasPorPais.decreto42_2025.manutencion[35] = 91.35''')
    
    doc.add_page_break()
    
    # =========================================================================
    # 7. CASOS ESPECIALES
    # =========================================================================
    doc.add_heading('7. Casos especiales', level=1)
    
    doc.add_heading('7.1. Viajes internacionales', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Los viajes internacionales requieren fechas de cruce de frontera. ')
    p.add_run('El sistema divide el viaje en hasta 3 tramos y calcula cada uno con sus propios precios.')
    
    p = doc.add_paragraph()
    p.add_run('Lógica de segmentación: ').bold = True
    p.add_run('calculoDesp.js → buildSegmentInputs()')
    
    doc.add_paragraph()
    
    doc.add_heading('7.2. Última noche ambigua', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Si el regreso es entre las 01:01 y las 06:59, la última noche está en ')
    p.add_run('"zona ambigua"').bold = True
    p.add_run('. El sistema:')
    
    bullets = [
        'Marca nochesAmbiguous = true',
        'Muestra un checkbox para que el usuario justifique la pernocta',
        'Si justifica, cuenta la noche; si no, no la cuenta'
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('Manejo del checkbox: ').bold = True
    p.add_run('salidaDesp.js → setupJustificarPernocta()')
    
    doc.add_paragraph()
    
    doc.add_heading('7.3. Ticket de cena (RD 462/2002)', level=2)
    
    p = doc.add_paragraph()
    p.add_run('En proyectos con normativa RD 462/2002, si el regreso es después de las 22:00 ')
    p.add_run('y se quiere contar la cena, ')
    p.add_run('es necesario aportar justificante de pago').bold = True
    p.add_run('.')
    
    p = doc.add_paragraph()
    p.add_run('Control de visibilidad: ').bold = True
    p.add_run('logicaDesp.js → shouldShowTicketCena()')
    
    doc.add_paragraph()
    
    doc.add_heading('7.4. Descuento por comidas de congreso', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Si el usuario indica que el congreso incluye comidas, ')
    p.add_run('se descuenta el 50% del precio de manutención por cada comida incluida ')
    p.add_run('del IRPF sujeto a retención.')
    
    p = doc.add_paragraph()
    p.add_run('Cálculo del descuento: ').bold = True
    p.add_run('formLogic.js → computeDescuentoManutencion()')
    
    p = doc.add_paragraph()
    p.add_run('Aplicación al IRPF: ').bold = True
    p.add_run('salidaDesp.js → calcIrpfAjustado()')
    
    doc.add_page_break()
    
    # =========================================================================
    # 8. PUNTOS CLAVE PARA MODIFICAR
    # =========================================================================
    doc.add_heading('8. Puntos clave para modificar la lógica', level=1)
    
    p = doc.add_paragraph()
    p.add_run('Esta sección indica dónde realizar cambios según el tipo de modificación necesaria.')
    
    doc.add_paragraph()
    
    # Tabla de modificaciones
    table = doc.add_table(rows=12, cols=3)
    table.style = 'Table Grid'
    
    hdr = table.rows[0].cells
    hdr[0].text = 'Si necesito...'
    hdr[1].text = 'Archivo'
    hdr[2].text = 'Función/Sección'
    for cell in hdr:
        set_cell_shading(cell, 'D9EAD3')
        cell.paragraphs[0].runs[0].bold = True
    
    modifications = [
        ('Cambiar precios de dietas', 'datos.json', 'dietasPorPais → decreto42_2025 o rd462_2002'),
        ('Cambiar tarifas de km', 'datos.json', 'kmTarifas'),
        ('Cambiar límites IRPF', 'datos.json', 'limitesIRPF'),
        ('Añadir un tipo de proyecto', 'datos.json', 'tiposProyecto + normativasPorTipoProyecto'),
        ('Cambiar horas de comida/cena', 'calculoDesp.js', 'Constantes HORA_COMIDA, HORA_CENA'),
        ('Cambiar reglas de manutención', 'calculoDesp.js', 'calcManutencionesSameDay(), calcManutencionesSeveralDays()'),
        ('Cambiar reglas de pernocta', 'calculoDesp.js', 'evalLastNightByHour(), calcNoches()'),
        ('Cambiar umbral residencia eventual', 'calculoDesp.js', 'isResidenciaEventual()'),
        ('Cambiar cálculo de IRPF', 'calculoDesp.js', 'calcIRPF()'),
        ('Modificar el formato visual', 'salidaDesp.js', 'templates.lineaManutencion(), etc.'),
        ('Cambiar validaciones de fechas', 'cogeDatosDesp.js', 'validarFechas()')
    ]
    
    for i, (need, archivo, funcion) in enumerate(modifications, 1):
        row = table.rows[i].cells
        row[0].text = need
        row[1].text = archivo
        row[2].text = funcion
    
    doc.add_paragraph()
    
    doc.add_heading('8.1. Ejemplo: Cambiar la hora de la comida', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Si quisiera cambiar la hora de la comida de 14:00 a 13:30:')
    
    p = doc.add_paragraph('1. Abrir ')
    p.add_run('calculoDesp.js').bold = True
    
    p = doc.add_paragraph('2. Buscar la constante:')
    add_code_block(doc, 'const HORA_COMIDA = 14 * 60;  // 14:00 = 840 minutos')
    
    p = doc.add_paragraph('3. Cambiar a:')
    add_code_block(doc, 'const HORA_COMIDA = 13 * 60 + 30;  // 13:30 = 810 minutos')
    
    doc.add_paragraph()
    
    doc.add_heading('8.2. Ejemplo: Añadir un nuevo precio por país', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Si se añade un nuevo país a las tablas:')
    
    p = doc.add_paragraph('1. Abrir ')
    p.add_run('datos.json').bold = True
    
    p = doc.add_paragraph('2. Añadir el país al array ')
    p.add_run('paises').bold = True
    p.add_run(' (al final, antes de "Resto del mundo")')
    
    p = doc.add_paragraph('3. Añadir el precio de alojamiento en ')
    p.add_run('decreto42_2025.alojamiento').bold = True
    p.add_run(' en la misma posición')
    
    p = doc.add_paragraph('4. Añadir el precio de manutención en ')
    p.add_run('decreto42_2025.manutencion').bold = True
    p.add_run(' en la misma posición')
    
    p = doc.add_paragraph('5. Repetir para ')
    p.add_run('rd462_2002').bold = True
    
    doc.add_paragraph()
    
    doc.add_heading('8.3. Recomendaciones generales', level=2)
    
    bullets = [
        'Hacer una copia de seguridad antes de modificar',
        'Probar con varios casos después del cambio',
        'Los cambios en datos.json son los más seguros (solo datos)',
        'Los cambios en calculoDesp.js son más delicados (lógica central)',
        'Usar la consola del navegador (F12) para depurar si algo falla'
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')
    
    # =========================================================================
    # GUARDAR
    # =========================================================================
    output_path = os.path.join(os.path.dirname(__file__), '..', 'docs', 'Logica_interna_calculo_desplazamientos.docx')
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f'Documento generado: {output_path}')
    return output_path

if __name__ == '__main__':
    create_document()
