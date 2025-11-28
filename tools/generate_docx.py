#!/usr/bin/env python3
"""
Script para generar `docs/Liquidacion_Desplazamientos.docx` con contenido explicativo.
Requisitos: pip install python-docx

Ejecución (PowerShell):
  python -m venv .venv
  .\.venv\Scripts\Activate.ps1
  pip install python-docx
  python tools\generate_docx.py

El script crea/actualiza `docs/Liquidacion_Desplazamientos.docx`.
"""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

text = {
    'title': 'Liquidación de un desplazamiento: Guía práctica paso a paso',
    'summary': 'Esta guía explica cómo calcular, con papel y calculadora, la liquidación económica de un desplazamiento profesional: dietas (manutención), noches/alojamiento, kilometraje, otros gastos y la retención por IRPF que corresponda. Describe qué operaciones realizar en función de los datos que el usuario introduce: tipo de proyecto, fechas y horas, país de destino, cruces de fronteras, importe real de alojamiento, justificantes de cena/pernocta, exclusión de manutención, kms y otros gastos.',
    'entries': [
        ('Tipo de proyecto', 'Indica la normativa aplicable (p. ej. "RD 462/2002" o un Decreto autonómico).'),
        ('Fechas y horas de salida y regreso', 'Fecha/hora de inicio y fecha/hora de fin del desplazamiento.'),
        ('País de destino y frontera (opcional)', 'Identifica si es nacional (España) o internacional (extranjero).'),
        ('Fechas de cruce de fronteras', 'Si el desplazamiento es internacional, la fecha (y hora) en que se sale y se vuelve a entrar al país de origen.'),
        ('Importe gastado en alojamiento (total)', 'Cantidad que el beneficiario pagó por hoteles en el desplazamiento (suma total o por tramo).'),
        ('Justificante de cena en la última noche', 'Si el beneficiario puede justificar que cenó la última noche del desplazamiento.'),
        ('Justificante de pernocta en la última noche', 'Si puede justificar que durmió la última noche (esto puede aumentar el cálculo de noches).'),
        ('Excluir manutención', 'Si el beneficiario no va a percibir gastos de manutención.'),
        ('Kilómetros recorridos y tipo de vehículo', 'Km totales y si es coche o moto (afecta tarifa por km).'),
        ('Otros gastos', 'Sumas adicionales justificadas (peajes, estacionamiento, etc.).'),
    ],
    'steps': [
        ('Paso 0: determinar la normativa aplicable', (
            'Identifique el tipo de proyecto: ciertas claves (por ejemplo G24, PEI, etc.) pueden estar mapeadas a la normativa "RD 462/2002" (en adelante "RD"). '
            'Otras claves aplican un Decreto autonómico/estatal (en adelante "Decreto"). '
            'La distinción importa porque las reglas para contar unidades de manutención (y cómo afecta un justificante de "ticket-cena") cambian levemente según la normativa.'
        )),
        ('Paso 1: segmentar el desplazamiento (si es internacional)', (
            'Nacional (sin cruces): trate el desplazamiento como un tramo desde la fecha/hora de ida hasta la fecha/hora de regreso.\n'
            'Internacional con cruces: divida el desplazamiento en hasta 3 tramos:\n'
            '  - Tramo A (ida interna): desde fechaIda@horaIda hasta cruceIda@08:00 (si la fecha de ida y la fecha de cruce inicial no coinciden).\n'
            '  - Tramo B (estancia en extranjero): desde cruceIda@08:00 hasta cruceVuelta@23:59.\n'
            '  - Tramo C (vuelta interna): desde cruceVuelta@00:00 hasta fechaRegreso@horaRegreso (si la fecha de cruce final y la fecha de regreso no coinciden).\n'
            'Nota: el tramo final es el que termina con la fecha de regreso real; éste es el único en el que debe comprobarse realmente el checkbox de "ticket-cena" cuando la normativa sea Decreto. '
            'Para Decreto, en los tramos no finales se asume que el beneficiario cenó.'
        )),
        ('Paso 2: calcular las unidades de manutención por tramo', (
            'Para cada tramo siga estas reglas:\n'
            ' - Si el tramo es "Same day": calcule la duración en horas; si no hay horas, manutención = 0. Si duración < 5 h puede considerarse 0.5 en casos concretos según normativa.\n'
            ' - Si cubre más de un día:\n'
            '    * Hora de salida: <14:00 => +1; 14:00-21:59 => +0.5; >=22:00 => +0.\n'
            '    * Días intermedios completos: cada día => +1.\n'
            '    * Hora de regreso: Decreto => >=22:00 => +1 (o >=14:00 => +0.5); RD => >=22:00 y ticket-cena => +1 (si no, >=14:00 => +0.5).\n'
            'Importante: cuando la normativa es Decreto y el tramo NO es el tramo final, se asume que el beneficiario cenó (por tanto la llegada cuenta como 1 unidad si es >=22:00).'
        )),
        ('Paso 3: convertir unidades en importe de manutención', (
            'Identifique el precio por manutención aplicable al país (tabla oficial).\n'
            'manutenciones_amount = unidades_total * precio_manutencion; redondear a 2 decimales.'
        )),
        ('Paso 4: calcular las noches y el alojamiento máximo', (
            'Noches base = diferencia de días entre medianoches.\n'
            'Noches contadas: si horaRegreso >=07:00 => contar todas; si horaRegreso <=01:00 => contar base -1; entre 01:00 y 07:00 => ambiguo (documentar o pedir justificante).\n'
            'Si se justifica pernocta última noche y la normativa lo permite, sumar +1 noche.\n'
            'Alojamiento máximo = noches_contadas * precio_noche. Reembolsar el menor entre gasto real y máximo.'
        )),
        ('Paso 5: kilometraje', 'kmAmount = km_recorridos * tarifa_km (según vehículo).'),
        ('Paso 6: otros gastos', 'Sume todos los gastos adicionales justificables.'),
        ('Paso 7: IRPF', (
            'Por día: calcular brutoOriginal = unidades_dia * precioManutencion; aplicar coeficiente (residencia eventual, p.ej. 0.8) => brutoToUse; comparar con límites (L1 para último día, L2 para intermedios) => sujeto = max(0, brutoToUse - exento).\n'
            'IRPF_total = suma(sujeto_por_dia). Si se excluye manutención, IRPF sujeto = 0.'
        )),
        ('Paso 8: sumar totales', 'Total = manutenciones + alojamiento_reembolsable + kmAmount + otros - IRPF_total.')
    ],
    'notes': (
        'Redondear siempre a dos decimales en operaciones monetarias intermedias.\n'
        'Conservar justificantes: facturas de alojamiento, tickets de peaje, justificantes de cenas/pernoctas.\n'
        'En casos ambiguos documentar la decisión.'
    )
    ,
    'input_schema_title': 'Especificación del objeto `input` para `calculateDesplazamiento`',
    'input_schema': [
        ('fechaIda', 'string', 'Obligatorio. Fecha de salida en formato dd/mm/aa o dd/mm/yyyy. Ej: "01/12/2025"'),
        ('horaIda', 'string', 'Obligatorio (salvo en modos segmentados). Hora de salida en formato hh:mm. Ej: "09:00"'),
        ('fechaRegreso', 'string', 'Obligatorio. Fecha de regreso en formato dd/mm/aa o dd/mm/yyyy. Ej: "03/12/2025"'),
        ('horaRegreso', 'string', 'Obligatorio (salvo en modos segmentados). Hora de regreso en formato hh:mm. Ej: "20:00"'),
        ('cruceIda', 'string', 'Opcional. Fecha del cruce de frontera de ida (solo para internacional). Formato dd/mm/aa.'),
        ('cruceVuelta', 'string', 'Opcional. Fecha del cruce de frontera de vuelta (solo para internacional).'),
        ('pais', 'string', 'Opcional. Nombre del país destino (p. ej. "Francia"). Se usa junto a `paisIndex` para elegir tarifas.'),
        ('paisIndex', 'number', 'Opcional pero preferible. Índice en la lista `dietasPorPais.paises` que identifica el país; 0 = España.'),
        ('km', 'number|string', 'Opcional. Kilómetros recorridos. Acepta número o cadena con separadores ("123,4" o "1.234").'),
        ('alojamiento', 'number|string', 'Opcional. Importe total de alojamiento en euros. Acepta número o cadena con formato monetario.'),
        ('ticketCena', 'boolean', 'Opcional. Indica si el beneficiario aporta justificante de la cena del último día.'),
        ('tipoProyecto', 'string', 'Opcional. Clave del tipo de proyecto; se busca en `normativasPorTipoProyecto.rd` para decidir reglas (RD vs Decreto).'),
        ('kmTarifa', 'number', 'Opcional. Tarifa por km para este cálculo (sobrescribe valores por defecto).'),
        ('residenciaEventual', 'boolean', 'Opcional. Si true aplica coeficiente de residencia eventual (p.ej. 0.8) para cálculo IRPF.'),
        ('flags', 'object', 'Opcional. Subobjeto con banderas de comportamiento: { excludeManutencion:boolean, justificarPernocta:boolean, excludeAlojamiento:boolean, _segmentMode:boolean }'),
        ('__nota__', 'string', 'El motor no espera otros campos; campos adicionales son ignorados. Para la generación de tramos, el wrapper puede construir inputs parciales con _segmentMode = true.')
    ],
    'input_example': (
        '{\n'
        '  "fechaIda": "01/12/2025",\n'
        '  "horaIda": "09:00",\n'
        '  "fechaRegreso": "03/12/2025",\n'
        '  "horaRegreso": "20:00",\n'
        '  "pais": "Francia",\n'
        '  "paisIndex": 1,\n'
        '  "km": "0",\n'
        '  "alojamiento": "150.00",\n'
        '  "ticketCena": false,\n'
        '  "tipoProyecto": "PROY-XYZ",\n'
        '  "kmTarifa": 0.26,\n'
        '  "residenciaEventual": false,\n'
        '  "flags": { "excludeManutencion": false, "justificarPernocta": false }\n'
        '}'
    ),
    'globals_used_title': 'Datos globales consultados por el motor',
    'globals_used': (
        'La función `calculateDesplazamiento` consulta `window.__sgtriDatos` para obtener tablas y parámetros necesarios:\n'
        '- `dietasPorPais`: estructura con `paises` (array) y subtablas por normativa (ej. `rd462_2002` y `decreto42_2025`) que contienen arrays `manutencion` y `alojamiento`.\n'
        '- `kmTarifas`: objeto con tarifas por tipo de vehículo (ej. `coche`, `moto`).\n'
        '- `limitesIRPF`: objeto con `esp` y `ext` arrays con límites exentos por día.\n'
        '- `normativasPorTipoProyecto`: objeto que incluye lista `rd` con las claves de `tipoProyecto` que aplican RD 462/2002.\n'
        'Si estas estructuras no existen en `window.__sgtriDatos`, el motor usa valores por defecto internos.'
    )
}


def create_docx(path):
    doc = Document()
    # Title
    h = doc.add_heading(level=1)
    run = h.add_run(text['title'])
    run.bold = True

    # Summary
    p = doc.add_paragraph()
    p.add_run('Resumen (para quien no conoce el tema)').bold = True
    doc.add_paragraph(text['summary'])

    doc.add_paragraph()
    doc.add_paragraph('Entradas necesarias').bold = True
    for k,v in text['entries']:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(k + ': ').bold = True
        p.add_run(v)

    doc.add_paragraph()
    for title, body in text['steps']:
        p = doc.add_paragraph()
        p.add_run(title).bold = True
        # split lines by \n to keep paragraphs
        for line in str(body).split('\n'):
            doc.add_paragraph(line)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Notas:').bold = True
    doc.add_paragraph(text['notes'])

    doc.save(path)

if __name__ == '__main__':
    import os
    outdir = os.path.join(os.path.dirname(__file__), '..', 'docs')
    os.makedirs(outdir, exist_ok=True)
    outpath = os.path.join(outdir, 'Liquidacion_Desplazamientos.docx')
    create_docx(outpath)
    print('Generado:', outpath)
