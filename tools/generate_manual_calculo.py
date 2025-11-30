#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Genera el documento Word "Cálculos de desplazamientos" - Manual de usuario
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_shading(cell, color):
    """Aplica color de fondo a una celda."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def create_manual():
    doc = Document()
    
    # Configurar márgenes
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # =========================================================================
    # TÍTULO PRINCIPAL
    # =========================================================================
    title = doc.add_heading('Cálculo de Desplazamientos', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Manual de cálculo para liquidaciones de dietas y desplazamientos')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.color.rgb = RGBColor(80, 80, 80)
    
    doc.add_paragraph()
    
    # =========================================================================
    # INTRODUCCIÓN
    # =========================================================================
    doc.add_heading('Introducción', level=1)
    
    intro = doc.add_paragraph()
    intro.add_run('Este manual explica paso a paso cómo calcular los importes correspondientes a un desplazamiento oficial. ')
    intro.add_run('El objetivo es que cualquier persona pueda verificar o calcular manualmente los importes de manutención, alojamiento y kilometraje aplicando las mismas reglas que utiliza el sistema.')
    
    doc.add_paragraph()
    
    # =========================================================================
    # SECCIÓN 1: CONCEPTOS BÁSICOS
    # =========================================================================
    doc.add_heading('1. Conceptos básicos', level=1)
    
    doc.add_heading('1.1. Componentes de una liquidación', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Una liquidación de desplazamiento puede incluir hasta cuatro conceptos:')
    
    bullet_items = [
        ('Manutención', 'Compensación por los gastos de comida durante el desplazamiento.'),
        ('Alojamiento', 'Reembolso de los gastos de hospedaje (con límites según normativa).'),
        ('Kilometraje', 'Compensación por el uso de vehículo particular.'),
        ('Otros gastos', 'Gastos adicionales elegibles (transporte, peajes, etc.).')
    ]
    
    for title_text, desc in bullet_items:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(title_text + ': ').bold = True
        p.add_run(desc)
    
    doc.add_paragraph()
    
    doc.add_heading('1.2. Normativas aplicables', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Existen dos normativas que determinan los importes:')
    
    # Tabla de normativas
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    
    # Encabezado
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Normativa'
    hdr_cells[1].text = 'Tipos de proyecto'
    for cell in hdr_cells:
        set_cell_shading(cell, 'D9EAD3')
        cell.paragraphs[0].runs[0].bold = True
    
    # Fila 1
    row1 = table.rows[1].cells
    row1[0].text = 'Decreto 42/2025 (Junta de Extremadura)'
    row1[1].text = 'Contratos UEx, Planes Complementarios, Otros Decreto'
    
    # Fila 2
    row2 = table.rows[2].cells
    row2[0].text = 'R.D. 462/2002 (Estatal)'
    row2[1].text = 'Ayudas GR24, Plan Estatal, Otros R.D. 462/2002'
    
    doc.add_paragraph()
    
    # =========================================================================
    # SECCIÓN 2: CÁLCULO DE MANUTENCIÓN
    # =========================================================================
    doc.add_heading('2. Cálculo de la manutención', level=1)
    
    doc.add_heading('2.1. El concepto de "manutención"', level=2)
    
    p = doc.add_paragraph()
    p.add_run('La manutención se mide en ')
    p.add_run('unidades').bold = True
    p.add_run('. Una unidad completa equivale a la dieta de un día entero. Media unidad (0,5) corresponde a una sola comida principal (almuerzo o cena).')
    
    doc.add_paragraph()
    
    doc.add_heading('2.2. Importes de manutención por normativa', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Para desplazamientos nacionales (España):')
    
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    
    hdr = table.rows[0].cells
    hdr[0].text = 'Normativa'
    hdr[1].text = 'Importe diario'
    hdr[2].text = 'Media manutención'
    for cell in hdr:
        set_cell_shading(cell, 'D9EAD3')
        cell.paragraphs[0].runs[0].bold = True
    
    row1 = table.rows[1].cells
    row1[0].text = 'Decreto 42/2025'
    row1[1].text = '53,34 €'
    row1[2].text = '26,67 €'
    
    row2 = table.rows[2].cells
    row2[0].text = 'R.D. 462/2002'
    row2[1].text = '37,40 €'
    row2[2].text = '18,70 €'
    
    p = doc.add_paragraph()
    p.add_run('Nota: ').bold = True
    p.add_run('Para desplazamientos internacionales, cada país tiene sus propios importes según las tablas oficiales.')
    
    doc.add_paragraph()
    
    doc.add_heading('2.3. Reglas de cálculo según la duración', level=2)
    
    # Caso A: Mismo día
    doc.add_heading('A) Desplazamiento de un solo día', level=3)
    
    p = doc.add_paragraph()
    p.add_run('Cuando la salida y el regreso son el mismo día, se aplican estas reglas:')
    
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    
    hdr = table.rows[0].cells
    hdr[0].text = 'Condición'
    hdr[1].text = 'Manutención'
    for cell in hdr:
        set_cell_shading(cell, 'FCE5CD')
        cell.paragraphs[0].runs[0].bold = True
    
    row1 = table.rows[1].cells
    row1[0].text = 'Sale antes de las 14:00 Y regresa después de las 16:00'
    row1[1].text = '0,5 (almuerzo)'
    
    row2 = table.rows[2].cells
    row2[0].text = 'Regresa después de las 22:00'
    row2[1].text = '+0,5 (cena)*'
    
    row3 = table.rows[3].cells
    row3[0].text = 'No cumple ninguna condición'
    row3[1].text = '0'
    
    p = doc.add_paragraph()
    p.add_run('* En proyectos R.D. 462/2002, la cena solo cuenta si se aporta justificante de pago.')
    p.runs[0].italic = True
    p.runs[0].font.size = Pt(10)
    
    doc.add_paragraph()
    
    # Caso B: Varios días
    doc.add_heading('B) Desplazamiento de varios días', level=3)
    
    p = doc.add_paragraph()
    p.add_run('Cuando el desplazamiento abarca varios días, se calcula por separado:')
    
    # Día de salida
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Día de salida:').bold = True
    
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    
    hdr = table.rows[0].cells
    hdr[0].text = 'Hora de salida'
    hdr[1].text = 'Manutención'
    for cell in hdr:
        set_cell_shading(cell, 'CFE2F3')
        cell.paragraphs[0].runs[0].bold = True
    
    table.rows[1].cells[0].text = 'Antes de las 14:00'
    table.rows[1].cells[1].text = '1 (completa)'
    table.rows[2].cells[0].text = 'Entre las 14:00 y las 22:00'
    table.rows[2].cells[1].text = '0,5 (cena)'
    table.rows[3].cells[0].text = 'Después de las 22:00'
    table.rows[3].cells[1].text = '0'
    
    doc.add_paragraph()
    
    # Días intermedios
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Días intermedios: ').bold = True
    p.add_run('1 manutención completa por cada día.')
    
    doc.add_paragraph()
    
    # Día de regreso
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Día de regreso:').bold = True
    
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    
    hdr = table.rows[0].cells
    hdr[0].text = 'Hora de regreso'
    hdr[1].text = 'Manutención'
    for cell in hdr:
        set_cell_shading(cell, 'CFE2F3')
        cell.paragraphs[0].runs[0].bold = True
    
    table.rows[1].cells[0].text = 'Después de las 22:00'
    table.rows[1].cells[1].text = '1 (completa)*'
    table.rows[2].cells[0].text = 'Entre las 14:00 y las 22:00'
    table.rows[2].cells[1].text = '0,5 (almuerzo)'
    table.rows[3].cells[0].text = 'Antes de las 14:00'
    table.rows[3].cells[1].text = '0'
    
    p = doc.add_paragraph()
    p.add_run('* En proyectos R.D. 462/2002, la cena solo cuenta si se aporta justificante.')
    p.runs[0].italic = True
    p.runs[0].font.size = Pt(10)
    
    doc.add_paragraph()
    
    doc.add_heading('2.4. Fórmula de cálculo', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Importe de manutención = Nº de unidades × Precio por unidad').bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Ejemplo
    doc.add_heading('Ejemplo práctico', level=3)
    
    example_box = doc.add_paragraph()
    example_box.add_run('Ejemplo: ').bold = True
    example_box.add_run('Desplazamiento del 15/01 a las 08:00 al 17/01 a las 19:00 (España, Decreto 42/2025)')
    
    bullets = [
        'Día 15 (salida a las 08:00): Sale antes de 14:00 → 1 manutención',
        'Día 16 (intermedio): → 1 manutención',
        'Día 17 (regreso a las 19:00): Regresa entre 14:00 y 22:00 → 0,5 manutención',
        'Total: 1 + 1 + 0,5 = 2,5 unidades',
        'Importe: 2,5 × 53,34 € = 133,35 €'
    ]
    
    for b in bullets:
        p = doc.add_paragraph(style='List Bullet')
        if '=' in b and '€' in b:
            p.add_run(b).bold = True
        else:
            p.add_run(b)
    
    doc.add_paragraph()
    
    # =========================================================================
    # SECCIÓN 3: CÁLCULO DE ALOJAMIENTO
    # =========================================================================
    doc.add_heading('3. Cálculo del alojamiento', level=1)
    
    doc.add_heading('3.1. Concepto de "noche"', level=2)
    
    p = doc.add_paragraph()
    p.add_run('El número de noches de alojamiento se calcula como los ')
    p.add_run('días naturales entre la fecha de salida y la fecha de regreso').bold = True
    p.add_run('. Por ejemplo:')
    
    bullets = [
        'Del 15/01 al 17/01 = 2 noches (las noches del 15 al 16, y del 16 al 17)',
        'Del 15/01 al 15/01 = 0 noches (mismo día)'
    ]
    
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')
    
    doc.add_paragraph()
    
    doc.add_heading('3.2. La última noche: regla de la hora de regreso', level=2)
    
    p = doc.add_paragraph()
    p.add_run('La última noche tiene reglas especiales según la hora de regreso:')
    
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    
    hdr = table.rows[0].cells
    hdr[0].text = 'Hora de regreso'
    hdr[1].text = '¿Cuenta la última noche?'
    for cell in hdr:
        set_cell_shading(cell, 'F4CCCC')
        cell.paragraphs[0].runs[0].bold = True
    
    table.rows[1].cells[0].text = 'A las 07:00 o después'
    table.rows[1].cells[1].text = 'SÍ cuenta'
    table.rows[2].cells[0].text = 'A la 01:00 o antes'
    table.rows[2].cells[1].text = 'NO cuenta'
    table.rows[3].cells[0].text = 'Entre las 01:01 y las 06:59'
    table.rows[3].cells[1].text = 'Zona ambigua (requiere justificación)'
    
    p = doc.add_paragraph()
    p.add_run('Explicación: ').bold = True
    p.add_run('Si regresa muy temprano (antes de las 01:00), se entiende que no ha pernoctado esa última noche. Si regresa a partir de las 07:00, sí ha pernoctado. En la franja intermedia, depende de si puede justificar la pernocta.')
    
    doc.add_paragraph()
    
    doc.add_heading('3.3. Importe máximo de alojamiento', level=2)
    
    p = doc.add_paragraph()
    p.add_run('El alojamiento tiene un ')
    p.add_run('importe máximo').bold = True
    p.add_run(' por noche según el país y la normativa. Se reembolsa el menor entre:')
    
    bullets = [
        'El gasto real justificado',
        'El límite máximo según tablas oficiales'
    ]
    
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('Límites para España:')
    
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    
    hdr = table.rows[0].cells
    hdr[0].text = 'Normativa'
    hdr[1].text = 'Máximo por noche'
    for cell in hdr:
        set_cell_shading(cell, 'D9EAD3')
        cell.paragraphs[0].runs[0].bold = True
    
    table.rows[1].cells[0].text = 'Decreto 42/2025'
    table.rows[1].cells[1].text = '102,56 €'
    table.rows[2].cells[0].text = 'R.D. 462/2002'
    table.rows[2].cells[1].text = '65,97 €'
    
    doc.add_paragraph()
    
    doc.add_heading('3.4. Fórmula de cálculo', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Alojamiento máximo = Nº de noches × Precio máximo por noche').bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    p.add_run('Alojamiento a liquidar = mínimo(Gasto real, Alojamiento máximo)').bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # =========================================================================
    # SECCIÓN 4: CÁLCULO DE KILOMETRAJE
    # =========================================================================
    doc.add_heading('4. Cálculo del kilometraje', level=1)
    
    p = doc.add_paragraph()
    p.add_run('El kilometraje compensa el uso de vehículo particular para el desplazamiento.')
    
    doc.add_heading('4.1. Tarifas por kilómetro', level=2)
    
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    
    hdr = table.rows[0].cells
    hdr[0].text = 'Tipo de vehículo'
    hdr[1].text = 'Tarifa por km'
    for cell in hdr:
        set_cell_shading(cell, 'D9EAD3')
        cell.paragraphs[0].runs[0].bold = True
    
    table.rows[1].cells[0].text = 'Coche'
    table.rows[1].cells[1].text = '0,26 €/km'
    table.rows[2].cells[0].text = 'Motocicleta'
    table.rows[2].cells[1].text = '0,106 €/km'
    
    doc.add_paragraph()
    
    doc.add_heading('4.2. Fórmula de cálculo', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Importe de kilometraje = Kilómetros × Tarifa por km').bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Ejemplo
    p = doc.add_paragraph()
    p.add_run('Ejemplo: ').bold = True
    p.add_run('240 km en coche = 240 × 0,26 € = 62,40 €')
    
    doc.add_paragraph()
    
    # =========================================================================
    # SECCIÓN 5: VIAJES INTERNACIONALES
    # =========================================================================
    doc.add_heading('5. Desplazamientos internacionales', level=1)
    
    p = doc.add_paragraph()
    p.add_run('Los viajes internacionales se dividen en ')
    p.add_run('tramos').bold = True
    p.add_run(' según el territorio donde se encuentre cada día:')
    
    bullets = [
        'Tramo España (ida): Desde la salida hasta el cruce de frontera de ida',
        'Tramo extranjero: Desde el cruce de frontera de ida hasta el cruce de frontera de vuelta',
        'Tramo España (vuelta): Desde el cruce de frontera de vuelta hasta el regreso'
    ]
    
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('Importante: ').bold = True
    p.add_run('Cada tramo utiliza los precios correspondientes a su territorio. El tramo en el extranjero usa las tablas del país de destino.')
    
    doc.add_paragraph()
    
    doc.add_heading('5.1. Ejemplo de viaje internacional', level=2)
    
    example = doc.add_paragraph()
    example.add_run('Viaje a Francia:').bold = True
    
    bullets = [
        'Salida: 10/03 a las 07:00 desde Badajoz',
        'Cruce frontera ida: 10/03 (por Irún)',
        'Cruce frontera vuelta: 14/03',
        'Regreso: 14/03 a las 20:00 a Badajoz'
    ]
    
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('Desglose:')
    
    bullets = [
        'Tramo España (ida): 10/03 → Precios España',
        'Tramo Francia: 10/03 al 14/03 → Precios Francia',
        'Tramo España (vuelta): 14/03 → Precios España'
    ]
    
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')
    
    doc.add_paragraph()
    
    # =========================================================================
    # SECCIÓN 6: RESIDENCIA EVENTUAL
    # =========================================================================
    doc.add_heading('6. Residencia eventual (estancias prolongadas)', level=1)
    
    p = doc.add_paragraph()
    p.add_run('Cuando un desplazamiento supera cierta duración, se considera ')
    p.add_run('"residencia eventual"').bold = True
    p.add_run(' y los importes de manutención y alojamiento se reducen al ')
    p.add_run('80%').bold = True
    p.add_run('.')
    
    doc.add_heading('6.1. Umbrales de duración', level=2)
    
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    
    hdr = table.rows[0].cells
    hdr[0].text = 'Destino'
    hdr[1].text = 'Se aplica reducción si supera'
    for cell in hdr:
        set_cell_shading(cell, 'F4CCCC')
        cell.paragraphs[0].runs[0].bold = True
    
    table.rows[1].cells[0].text = 'España'
    table.rows[1].cells[1].text = '1 mes'
    table.rows[2].cells[0].text = 'Extranjero'
    table.rows[2].cells[1].text = '3 meses'
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('Ejemplo: ').bold = True
    p.add_run('Un desplazamiento en España del 01/02 al 05/03 (más de 1 mes) tendría los importes de manutención y alojamiento multiplicados por 0,8.')
    
    doc.add_paragraph()
    
    # =========================================================================
    # SECCIÓN 7: IRPF
    # =========================================================================
    doc.add_heading('7. Retención de IRPF', level=1)
    
    p = doc.add_paragraph()
    p.add_run('La manutención está sujeta a retención de IRPF por la parte que exceda los ')
    p.add_run('límites exentos').bold = True
    p.add_run('.')
    
    doc.add_heading('7.1. Límites exentos diarios', level=2)
    
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    
    hdr = table.rows[0].cells
    hdr[0].text = 'Territorio'
    hdr[1].text = 'Sin pernocta'
    hdr[2].text = 'Con pernocta'
    for cell in hdr:
        set_cell_shading(cell, 'CFE2F3')
        cell.paragraphs[0].runs[0].bold = True
    
    table.rows[1].cells[0].text = 'España'
    table.rows[1].cells[1].text = '26,67 €'
    table.rows[1].cells[2].text = '53,34 €'
    
    table.rows[2].cells[0].text = 'Extranjero'
    table.rows[2].cells[1].text = '48,08 €'
    table.rows[2].cells[2].text = '91,35 €'
    
    doc.add_paragraph()
    
    doc.add_heading('7.2. Cálculo del importe sujeto a IRPF', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Para cada día del desplazamiento:')
    
    p = doc.add_paragraph()
    p.add_run('IRPF del día = máximo(0, Manutención del día − Límite exento)').bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    p.add_run('Se usa el límite "con pernocta" para todos los días excepto el último día del desplazamiento, que usa el límite "sin pernocta".')
    
    doc.add_paragraph()
    
    # =========================================================================
    # SECCIÓN 8: RESUMEN Y FÓRMULA FINAL
    # =========================================================================
    doc.add_heading('8. Cálculo del total', level=1)
    
    p = doc.add_paragraph()
    p.add_run('El importe total de la liquidación es:')
    
    p = doc.add_paragraph()
    formula = p.add_run('TOTAL = Manutención + Alojamiento + Kilometraje + Otros gastos')
    formula.bold = True
    formula.font.size = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('Donde:')
    
    bullets = [
        'Manutención = Unidades calculadas × Precio según normativa y país',
        'Alojamiento = mínimo(Gasto real, Nº noches × Precio máximo)',
        'Kilometraje = Kilómetros × Tarifa del vehículo',
        'Otros gastos = Suma de gastos adicionales elegibles'
    ]
    
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')
    
    doc.add_paragraph()
    
    # =========================================================================
    # ANEXO: TABLAS DE PRECIOS
    # =========================================================================
    doc.add_page_break()
    doc.add_heading('Anexo: Precios de referencia (España)', level=1)
    
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    
    hdr = table.rows[0].cells
    hdr[0].text = 'Concepto'
    hdr[1].text = 'Decreto 42/2025'
    hdr[2].text = 'R.D. 462/2002'
    for cell in hdr:
        set_cell_shading(cell, 'D9EAD3')
        cell.paragraphs[0].runs[0].bold = True
    
    table.rows[1].cells[0].text = 'Manutención diaria'
    table.rows[1].cells[1].text = '53,34 €'
    table.rows[1].cells[2].text = '37,40 €'
    
    table.rows[2].cells[0].text = 'Alojamiento máximo/noche'
    table.rows[2].cells[1].text = '102,56 €'
    table.rows[2].cells[2].text = '65,97 €'
    
    table.rows[3].cells[0].text = 'Kilometraje (coche)'
    table.rows[3].cells[1].text = '0,26 €/km'
    table.rows[3].cells[2].text = '0,26 €/km'
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('Nota: ').bold = True
    p.add_run('Para destinos internacionales, consulte las tablas oficiales de cada normativa que incluyen importes específicos para cada país.')
    
    # =========================================================================
    # GUARDAR DOCUMENTO
    # =========================================================================
    output_path = os.path.join(os.path.dirname(__file__), '..', 'docs', 'Calculo_de_desplazamientos.docx')
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f'Documento generado: {output_path}')
    return output_path

if __name__ == '__main__':
    create_manual()
